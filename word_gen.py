import docx
import os
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import barcode_excel_gen
from web_parse import web_parse


def createFolder(directory):  # Создание папки docx
    if os.path.exists(directory+'/docx'):
        pass
    else:
        os.makedirs(directory+'/docx')


def dirParse():  # Поиск изображений баркодов
    try:
        os.chdir('./folder')
    except:
        return 'Folder is missing'
    names = os.listdir()
    os.chdir('..')
    return names


def gen_word(full_path, country, firm, year):
    directory, path, filename, ext = barcode_excel_gen.pathSplit(full_path)
    os.chdir(directory)
    createFolder(directory)

    # Парсинг отчета
    barcodes = dirParse()
    art_pos, size, code, types, color, art_wb = barcode_excel_gen.parse(full_path)
    if barcodes == 'Folder is missing':
        return 'Folder is missing'
    elif barcodes == []:
        return 'Barcodes are missing'
    else:
        # Создание документов по каждому баркоду
        for image_name in barcodes:
            # Текущее описание
            im1, im2 = image_name.split(sep='_')
            image, im3 = im2.split(sep='.')
            num_code = code.index(image)
            cur_art_pos = art_pos[num_code]
            cur_color = color[num_code]
            cur_size = size[num_code]
            cur_type = types[num_code]
            filename, ext = os.path.splitext(image_name)
            doc = docx.Document()

            # Парсинг вебсайта для получения состава
            cur_consist = web_parse(art_wb[num_code])

            # Размер и разметка страницы
            section = doc.sections[0]
            section.page_height = Cm(5.8)
            section.page_width = Cm(6.0)
            section.left_margin = Cm(0.5)
            section.right_margin = Cm(0.5)
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.1)

            # Типичное форматирование, применяющееся ко всему тексту
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(5)

            # добавим изображение
            p1 = doc.add_paragraph()
            p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p1.paragraph_format.space_after = Pt(0)
            p1.style = doc.styles['Normal']
            r1 = p1.add_run()
            r1.add_picture('./folder/' + image_name, width=Cm(5.0), height=Cm(2.0))

            # добавим описание
            p2 = doc.add_paragraph('Наименование: ' + str(cur_type) + '\n'
                        'Артикул: ' + str(cur_art_pos) + '\n'
                        'Цвет: ' + str(cur_color) + '\n'
                        'Размер: ' + str(cur_size) + '\n'
                        'Состав: ' + str(cur_consist) + '\n'
                        'Страна производства: ' + str(country) + '\n'
                        'Производитель: ' + str(firm) + '\n'
                        'Дата изготовления: ' + str(year))
            p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
            p2.style = doc.styles['Normal']

            # Добавим EAC
            p3 = doc.add_paragraph('TP TC 017/2011\t\t\t')
            p3.paragraph_format.space_after = Pt(0)
            p3.style = doc.styles['Normal']
            r3 = p3.add_run()
            r3.add_picture('eac.png', width=Cm(0.8), height=Cm(0.8), )

            # В случае обнаружения нового цвета, помечаем документ приставкой "Redo"
            if cur_color == '-':
                doc.save('./docx/' + filename + '_Redo.docx')
            else:
                doc.save('./docx/' + filename + '.docx')
